from mapper import map_text_to_runs, merge_diacritics, compute_diacritics_stats, save_highlighted_copy


def test_map_single_run():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph("A inceput sa ii taie un fir de par.")
    run = p.runs[0]

    runs_info = [{"run": run, "text": run.text}]
    map_text_to_runs(
        "A inceput sa ii taie un fir de par.",
        "a început să îi taie un fir de păr",
        runs_info,
    )
    assert run.text == "A început să îi taie un fir de păr."


def test_map_multiple_runs():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("A inceput ")
    r2 = p.add_run("sa ii taie ")
    r3 = p.add_run("un fir de par.")

    original = "A inceput sa ii taie un fir de par."
    corrected = "A început să îi taie un fir de păr."
    runs_info = [
        {"run": r1, "text": "A inceput "},
        {"run": r2, "text": "sa ii taie "},
        {"run": r3, "text": "un fir de par."},
    ]
    map_text_to_runs(original, corrected, runs_info)
    assert "început" in r1.text
    assert "să" in r2.text
    assert "păr" in r3.text


def test_map_empty_runs():
    runs_info = []
    map_text_to_runs("text", "correction", runs_info)
    # Should not raise


def test_merge_preserves_capitalization():
    assert merge_diacritics("A inceput", "a început") == "A început"


def test_merge_all_caps():
    assert merge_diacritics("INCEPUT", "început") == "ÎNCEPUT"


def test_merge_title_case():
    assert merge_diacritics("Inceput", "început") == "Început"


def test_merge_punctuation():
    assert merge_diacritics("par.", "păr") == "păr."


def test_merge_partial_diacritics():
    assert merge_diacritics("să inceput", "să început") == "să început"


def test_merge_no_match_fallback():
    # different word — model bug, keep original
    assert merge_diacritics("original", "total_diferit") == "original"


def test_merge_multiple_words():
    original = "A inceput sa ii taie un fir de par."
    corrected = "a început să îi taie un fir de păr"
    result = merge_diacritics(original, corrected)
    assert result == "A început să îi taie un fir de păr."


def test_map_preserves_whitespace():
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("A  ")  # two spaces
    r2 = p.add_run("inceput ")  # trailing space
    r3 = p.add_run("sa  ii  taie un fir de par.")  # double spaces

    original = "A  inceput sa  ii  taie un fir de par."
    corrected = "a  început să  ii  taie un fir de păr"
    runs_info = [
        {"run": r1, "text": "A  "},
        {"run": r2, "text": "inceput "},
        {"run": r3, "text": "sa  ii  taie un fir de par."},
    ]
    map_text_to_runs(original, corrected, runs_info)
    # "ii" doesn't change in this context; "sa"→"să", "par."→"păr."
    assert r1.text == "A  "  # two spaces preserved
    assert "început" in r2.text
    assert r3.text == "să  ii  taie un fir de păr."  # spaces preserved, diacritice aplicate


def test_map_preserves_formatting():
    """Runs keep their bold/italic/etc. after mapping."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Primul ")
    r1.bold = True
    r1.font.size = Pt(14)
    r2 = p.add_run("Al doilea ")
    r2.italic = True

    original = "Primul Al doilea "
    corrected = "Primul Al doilea "
    runs_info = [
        {"run": r1, "text": "Primul "},
        {"run": r2, "text": "Al doilea "},
    ]
    map_text_to_runs(original, corrected, runs_info)
    assert r1.bold is True
    assert r1.font.size == Pt(14)
    assert r2.italic is True


def test_compute_stats():
    orig = "A inceput sa ii taie un fir de par"
    merged = "A început să îi taie un fir de păr"
    stats = compute_diacritics_stats(orig, merged)
    assert stats["total"] == 9
    assert stats["modified"] == 4
    assert stats["potential"] == 3  # A, taie, fir — unchanged but can have diacritics


def test_compute_stats_with_potential():
    orig = "carte par masina"
    merged = "carte par mașină"
    stats = compute_diacritics_stats(orig, merged)
    assert stats["total"] == 3
    assert stats["modified"] == 1
    assert stats["potential"] == 2  # carte, par — unchanged but can


def test_compute_stats_no_changes():
    orig = "carte stat mare"
    merged = "carte stat mare"
    stats = compute_diacritics_stats(orig, merged)
    assert stats["total"] == 3
    assert stats["modified"] == 0
    assert stats["potential"] == 3  # carte, stat, mare all have a/s/t


def test_compute_stats_already_diacritic():
    orig = "început păr"
    merged = "început păr"
    stats = compute_diacritics_stats(orig, merged)
    assert stats["total"] == 2
    assert stats["modified"] == 0
    assert stats["potential"] == 0  # already have diacritics, not counted


def test_save_highlighted(tmp_path):
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("A inceput sa")
    src = tmp_path / "src.docx"
    doc.save(str(src))

    data = [
        {
            "orig_words": ["A", "inceput", "sa"],
            "merged_words": ["A", "început", "să"],
        }
    ]
    out = tmp_path / "hl.docx"
    save_highlighted_copy(str(src), str(out), data)

    result = Document(str(out))
    runs = list(result.paragraphs[0].runs)
    texts = [r.text for r in runs]
    highlights = [r.font.highlight_color for r in runs]
    assert "început" in texts or "început " in texts
    assert "să" in texts or "să " in texts
    assert any(h == WD_COLOR_INDEX.YELLOW for h in highlights if h is not None)


def test_merge_ignores_url():
    assert merge_diacritics("https://example.com", "https://example.com") == "https://example.com"
    assert merge_diacritics("http://site.ro/ceva", "http://site.ro/ceva") == "http://site.ro/ceva"


def test_merge_ignores_email():
    assert merge_diacritics("contact@domeniu.ro", "contact@domeniu.ro modificat") == "contact@domeniu.ro"


def test_merge_ignores_url_in_text():
    orig = "Vezi pe https://example.com pentru detalii"
    corr = "vezi pe https://example.com pentru detalii"
    result = merge_diacritics(orig, corr)
    assert "https://example.com" in result
    assert result.startswith("Vezi")  # capitalizare păstrată
