import subprocess
import sys
from pathlib import Path


def test_e2e_diacritics_restoration(tmp_path):
    from docx import Document

    # Create a test document with mixed formatting
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("A inceput ").bold = True
    p.add_run("sa ii taie ").italic = True
    p.add_run("un fir de par.")

    input_path = tmp_path / "input.docx"
    doc.save(str(input_path))
    output_path = tmp_path / "output.docx"

    # Run the CLI using same venv Python
    result = subprocess.run(
        [sys.executable, "corecteaza_diacritice.py", str(input_path), "-o", str(output_path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        timeout=300,
        cwd=Path(__file__).resolve().parent.parent,
    )
    assert result.returncode == 0, f"CLI failed:\nstdout: {result.stdout}\nstderr: {result.stderr}"

    # Verify output
    result_doc = Document(str(output_path))
    assert len(result_doc.paragraphs) == 1
    full_text = result_doc.paragraphs[0].text
    assert "început" in full_text, f"Missing 'început' in: {full_text}"
    assert "să" in full_text, f"Missing 'să' in: {full_text}"
    assert "păr" in full_text, f"Missing 'păr' in: {full_text}"

    # Verify formatting preserved
    runs = result_doc.paragraphs[0].runs
    for run in runs:
        if "început" in run.text:
            assert run.bold is True, f"Expected bold on: {run.text}"
        if "să" in run.text:
            assert run.italic is True, f"Expected italic on: {run.text}"
