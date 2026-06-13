from extractor import extract_paragraphs


def test_extract_single_paragraph(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("A inceput sa ii taie un fir de par.")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    paragraphs, doc_obj = extract_paragraphs(str(path))
    assert len(paragraphs) == 1
    assert paragraphs[0]["text"] == "A inceput sa ii taie un fir de par."
    assert len(paragraphs[0]["runs"]) == 1


def test_extract_paragraph_with_multiple_runs(tmp_path):
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("A inceput ")
    r1.bold = True
    r2 = p.add_run("sa ii taie ")
    r2.italic = True
    r3 = p.add_run("un fir de par.")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    paragraphs, doc_obj = extract_paragraphs(str(path))
    assert len(paragraphs) == 1
    assert paragraphs[0]["runs"][0]["text"] == "A inceput "
    assert paragraphs[0]["runs"][1]["text"] == "sa ii taie "
    assert paragraphs[0]["runs"][2]["text"] == "un fir de par."


def test_extract_skips_empty_paragraphs(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("Text real")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    paragraphs, doc_obj = extract_paragraphs(str(path))
    assert len(paragraphs) == 1
    assert paragraphs[0]["text"] == "Text real"
