import re


_DIACRITICS_MAP = {
    "ă": "a", "Ă": "A",
    "â": "a", "Â": "A",
    "î": "i", "Î": "I",
    "ș": "s", "Ș": "S",
    "ț": "t", "Ț": "T",
}


def _strip_diacritics(word: str) -> str:
    return "".join(_DIACRITICS_MAP.get(c, c) for c in word)


def merge_diacritics(original: str, corrected: str) -> str:
    parts = re.split(r"(\s+)", original)
    corr_words = corrected.split()
    out = []
    wi = 0
    for p in parts:
        if p.isspace():
            out.append(p)
        else:
            cw = corr_words[wi] if wi < len(corr_words) else p
            out.append(_transfer(p, cw))
            wi += 1
    return "".join(out)


_URL_EMAIL_RE = re.compile(
    r'^(?:https?://|ftp://|www\.)[^\s]*$|^[^\s@]+@[^\s@]+\.[^\s@]+$',
    re.IGNORECASE,
)


def _transfer(orig_word: str, corr_word: str) -> str:
    if _URL_EMAIL_RE.match(orig_word.rstrip(".,!?;:\"'-")):
        return orig_word
    o_clean = _strip_diacritics(orig_word).rstrip(".,!?;:\"'-").lower()
    c_clean = _strip_diacritics(corr_word).rstrip(".,!?;:\"'-").lower()
    if o_clean != c_clean:
        return orig_word

    orig_chars = list(orig_word)
    ci = 0
    for oi, oc in enumerate(orig_chars):
        if ci >= len(corr_word):
            break

        cc = corr_word[ci]
        if oc.isalpha() and cc.isalpha() and _strip_diacritics(oc).lower() == _strip_diacritics(cc).lower():
            orig_chars[oi] = cc.upper() if oc.isupper() else cc.lower()
            ci += 1
        elif oc.isalpha() and cc.isalpha():
            ci += 1
        else:
            if oc.isalpha():
                ci += 1
    return "".join(orig_chars)


def compute_diacritics_stats(original: str, merged: str) -> dict:
    orig_words = original.split()
    merged_words = merged.split()

    total = len(orig_words)
    modified = 0
    could_have = 0

    _CAN_DIACRITIC = set("aAiIsStt")

    for ow, mw in zip(orig_words, merged_words):
        if ow != mw:
            modified += 1
        elif any(c in _CAN_DIACRITIC for c in ow):
            if not any(c in "ăâîșțĂÂÎȘȚ" for c in ow):
                could_have += 1

    return {
        "total": total,
        "potential": could_have,
        "modified": modified,
    }


def save_highlighted_copy(
    input_path: str,
    output_path: str,
    paragraphs_data: list[dict],
) -> None:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document(input_path)
    pi = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if pi >= len(paragraphs_data):
            break

        data = paragraphs_data[pi]
        pi += 1
        orig_words = data["orig_words"]
        merged_words = data["merged_words"]

        modified = [i for i, (o, m) in enumerate(zip(orig_words, merged_words)) if o != m]
        if not modified:
            continue

        orig_runs = list(para.runs)
        if not orig_runs:
            continue

        tmpl = orig_runs[0]
        for run in orig_runs:
            run._element.getparent().remove(run._element)

        for wi, word in enumerate(merged_words):
            new_run = para.add_run(word)
            new_run.bold = tmpl.bold
            new_run.italic = tmpl.italic
            new_run.underline = tmpl.underline
            new_run.font.size = tmpl.font.size
            new_run.font.name = tmpl.font.name
            if wi in modified:
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if wi < len(merged_words) - 1:
                para.add_run(" ")

    doc.save(output_path)


def map_text_to_runs(
    original_text: str,
    corrected_text: str,
    runs_info: list[dict],
) -> None:
    if not runs_info:
        return

    merged = merge_diacritics(original_text, corrected_text)

    if len(runs_info) == 1:
        runs_info[0]["run"].text = merged
        return

    pos = 0
    for info in runs_info:
        run_text = info["text"]
        start = original_text.find(run_text, pos)
        if start == -1:
            pos += max(len(run_text), 1)
            continue
        end = start + len(run_text)
        info["run"].text = merged[start:end]
        pos = end
